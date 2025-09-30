#!/usr/bin/env python3
"""
Mail Merge Script for Real Word Documents
Works with existing .docx templates in the templates folder
"""

import os
import pandas as pd
from datetime import datetime
import re

try:
    from docx import Document
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False
    print("Warning: python-docx not installed. Install with: pip install python-docx")

class WordMailMergeProcessor:
    def __init__(self, data_file="data/test_data.csv", template_dir="templates", output_dir="generated_docs"):
        self.data_file = data_file
        self.template_dir = template_dir
        self.output_dir = output_dir
        self.data = None
        self.load_data()
        
        # Create output directory if it doesn't exist
        if not os.path.exists(self.output_dir):
            os.makedirs(self.output_dir)
            print(f"Created output directory: {self.output_dir}")
    
    def load_data(self):
        """Load data from CSV file"""
        try:
            if os.path.exists(self.data_file):
                self.data = pd.read_csv(self.data_file)
                print(f"✅ Loaded {len(self.data)} records from {self.data_file}")
                print(f"Columns: {list(self.data.columns)}")
            else:
                print(f"❌ Data file not found: {self.data_file}")
                self.data = None
        except Exception as e:
            print(f"❌ Error loading data: {e}")
            self.data = None
    
    def safe_filename(self, name):
        """Create a safe filename from employee name"""
        # Remove or replace unsafe characters
        safe_name = re.sub(r'[<>:"/\\|?*]', '_', str(name))
        safe_name = safe_name.replace(' ', '_')
        return safe_name
    
    def replace_placeholders_in_paragraph(self, paragraph, row_data):
        """Replace placeholders in a paragraph"""
        text = paragraph.text
        
        # Replace all {{field}} placeholders
        for column, value in row_data.items():
            placeholder = f"{{{{{column}}}}}"
            if placeholder in text:
                text = text.replace(placeholder, str(value))
        
        # Clear the paragraph and add the updated text
        paragraph.clear()
        paragraph.add_run(text)
    
    def replace_placeholders_in_table(self, table, row_data):
        """Replace placeholders in table cells"""
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    self.replace_placeholders_in_paragraph(paragraph, row_data)
    
    def process_word_template(self, template_path, row_data, output_filename=None):
        """Process a Word document template with employee data"""
        if not DOCX_AVAILABLE:
            print("❌ python-docx not available. Cannot process Word documents.")
            return None
        
        try:
            # Load the template document
            doc = Document(template_path)
            
            # Replace placeholders in paragraphs
            for paragraph in doc.paragraphs:
                self.replace_placeholders_in_paragraph(paragraph, row_data)
            
            # Replace placeholders in tables
            for table in doc.tables:
                self.replace_placeholders_in_table(table, row_data)
            
            # Replace placeholders in headers and footers
            for section in doc.sections:
                # Header
                if section.header:
                    for paragraph in section.header.paragraphs:
                        self.replace_placeholders_in_paragraph(paragraph, row_data)
                # Footer
                if section.footer:
                    for paragraph in section.footer.paragraphs:
                        self.replace_placeholders_in_paragraph(paragraph, row_data)
            
            # Generate output filename
            if output_filename is None:
                template_name = os.path.splitext(os.path.basename(template_path))[0]
                safe_employee_name = self.safe_filename(row_data.get('Name', 'Unknown'))
                output_filename = f"{template_name}_{safe_employee_name}.docx"
            
            output_path = os.path.join(self.output_dir, output_filename)
            
            # Save the processed document
            doc.save(output_path)
            print(f"✅ Generated: {output_filename}")
            
            return output_path
            
        except Exception as e:
            print(f"❌ Error processing template {template_path}: {e}")
            return None
    
    def process_all_templates(self):
        """Process all Word template files for all employees"""
        if self.data is None:
            print("❌ No data loaded. Cannot process templates.")
            return []
        
        # Create template directory if it doesn't exist
        if not os.path.exists(self.template_dir):
            os.makedirs(self.template_dir)
            print(f"Created template directory: {self.template_dir}")
            return []
        
        # Find all Word template files
        template_files = [f for f in os.listdir(self.template_dir) 
                         if f.endswith('.docx') and not f.startswith('~')]
        
        if not template_files:
            print(f"❌ No .docx template files found in {self.template_dir}")
            return []
        
        print(f"📄 Found {len(template_files)} template files:")
        for template in template_files:
            print(f"   - {template}")
        
        results = []
        total_processed = 0
        
        # Process each template for each employee
        for template_file in template_files:
            template_path = os.path.join(self.template_dir, template_file)
            print(f"\n🔄 Processing template: {template_file}")
            
            template_results = 0
            for index, row in self.data.iterrows():
                row_data = row.to_dict()
                
                output_path = self.process_word_template(template_path, row_data)
                
                if output_path:
                    results.append({
                        'template': template_file,
                        'employee': row_data.get('Name', f'Employee_{index}'),
                        'output': output_path,
                        'status': 'success'
                    })
                    template_results += 1
                    total_processed += 1
                else:
                    results.append({
                        'template': template_file,
                        'employee': row_data.get('Name', f'Employee_{index}'),
                        'output': None,
                        'status': 'failed'
                    })
            
            print(f"   ✅ Generated {template_results} documents from {template_file}")
        
        print(f"\n🎉 Processing complete! Generated {total_processed} documents total.")
        return results
    
    def generate_summary_report(self, results):
        """Generate a summary report of the processing results"""
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        summary_filename = f"processing_summary_{timestamp}.txt"
        summary_path = os.path.join(self.output_dir, summary_filename)
        
        try:
            with open(summary_path, 'w', encoding='utf-8') as f:
                f.write("=" * 60 + "\n")
                f.write("WORD DOCUMENT MAIL MERGE PROCESSING SUMMARY\n")
                f.write("=" * 60 + "\n")
                f.write(f"Generated: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}\n")
                f.write(f"Data Source: {self.data_file}\n")
                f.write(f"Template Directory: {self.template_dir}\n")
                f.write(f"Output Directory: {self.output_dir}\n")
                f.write("\n")
                
                successful = [r for r in results if r['status'] == 'success']
                failed = [r for r in results if r['status'] == 'failed']
                
                f.write(f"Total Documents Generated: {len(successful)}\n")
                f.write(f"Failed Processes: {len(failed)}\n")
                f.write("\n")
                
                if successful:
                    f.write("SUCCESSFUL GENERATIONS:\n")
                    f.write("-" * 30 + "\n")
                    for result in successful:
                        f.write(f"Template: {result['template']}\n")
                        f.write(f"Employee: {result['employee']}\n")
                        f.write(f"Output: {os.path.basename(result['output'])}\n")
                        f.write("\n")
                
                if failed:
                    f.write("FAILED GENERATIONS:\n")
                    f.write("-" * 30 + "\n")
                    for result in failed:
                        f.write(f"Template: {result['template']}\n")
                        f.write(f"Employee: {result['employee']}\n")
                        f.write("Error: Processing failed\n")
                        f.write("\n")
            
            print(f"📊 Summary report saved: {summary_filename}")
            return summary_path
            
        except Exception as e:
            print(f"❌ Error generating summary: {e}")
            return None

def main():
    """Main function to run the mail merge process"""
    print("🚀 Starting Word Document Mail Merge Process")
    print("=" * 50)
    
    # Check if python-docx is available
    if not DOCX_AVAILABLE:
        print("❌ python-docx library is required!")
        print("Install it with: pip install python-docx")
        return
    
    # Initialize processor
    processor = WordMailMergeProcessor(
        data_file="data/test_data.csv",
        template_dir="templates", 
        output_dir="generated_docs"
    )
    
    # Process all templates
    results = processor.process_all_templates()
    
    if results:
        # Generate summary report
        processor.generate_summary_report(results)
        print(f"\n📁 Check the '{processor.output_dir}' folder for generated documents!")
    else:
        print("\n❌ No documents were generated.")
        print("Please check:")
        print("- Data file exists and has valid data")
        print("- Template directory contains .docx files")
        print("- Templates have proper {{FieldName}} placeholders")

if __name__ == "__main__":
    main()